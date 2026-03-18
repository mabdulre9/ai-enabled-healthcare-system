from flask import Flask, render_template, request, jsonify, send_from_directory, send_file
import json
import os
from datetime import datetime
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

app = Flask(__name__)

# Data file path
DATA_FILE = 'data/patients.json'
SETTINGS_FILE = 'data/settings.json'


# Clinic information for reports
CLINIC_NAME = 'Neighborhood Clinic'
CLINIC_ADDRESS = '123 Medical Center Dr, City, State 12345'
CLINIC_PHONE = '(555) 123-4567'
CLINIC_EMAIL = 'contact@neighborhoodclinic.com'

# Email configuration (optional - for sending reports)
EMAIL_ENABLED = False  # Set to True after configuring SMTP
SMTP_SERVER = 'smtp.gmail.com'
SMTP_PORT = 587
EMAIL_ADDRESS = 'your-clinic-email@gmail.com'
EMAIL_PASSWORD = 'your-app-password'  # Use Gmail app-specific password

# Ollama AI Configuration (100% Offline & Private)
OLLAMA_SETTINGS = {
    'ollama_model': 'qwen2.5:4b',
    'ollama_host': 'http://localhost:11434',
    'system_prompt': 'You are a clinical AI assistant. Provide evidence-based medical insights.',
    'temperature': 0.7,
    'max_tokens': 2048
}



# Ensure data directory exists
os.makedirs('data', exist_ok=True)

# Initialize empty patients file if it doesn't exist
if not os.path.exists(DATA_FILE):
    with open(DATA_FILE, 'w') as f:
        json.dump([], f)

# Initialize settings file if it doesn't exist
if not os.path.exists(SETTINGS_FILE):
    with open(SETTINGS_FILE, 'w') as f:
        json.dump(OLLAMA_SETTINGS, f, indent=2)

def load_settings():
    """Load LLM settings from file"""
    try:
        with open(SETTINGS_FILE, 'r') as f:
            return json.load(f)
    except:
        return OLLAMA_SETTINGS.copy()

def save_settings(settings):
    """Save LLM settings to file"""
    with open(SETTINGS_FILE, 'w') as f:
        json.dump(settings, f, indent=2)
    # Update global settings
    # Settings updated

def load_patients():
    """Load all patients from JSON file"""
    try:
        with open(DATA_FILE, 'r') as f:
            return json.load(f)
    except:
        return []

def save_patients(patients):
    """Save patients to JSON file"""
    with open(DATA_FILE, 'w') as f:
        json.dump(patients, f, indent=2)

def generate_patient_id():
    """Generate next patient ID"""
    patients = load_patients()
    if not patients:
        return 'PATIENT-001'
    
    max_id = 0
    for patient in patients:
        num = int(patient['patientId'].split('-')[1])
        if num > max_id:
            max_id = num
    
    return f'PATIENT-{str(max_id + 1).zfill(3)}'

def generate_visit_id(patient_id):
    """Generate visit ID for a patient"""
    patients = load_patients()
    patient = next((p for p in patients if p['patientId'] == patient_id), None)
    
    if not patient or not patient.get('visits'):
        return f'{patient_id}-VISIT-001'
    
    max_id = len(patient['visits'])
    return f'{patient_id}-VISIT-{str(max_id + 1).zfill(3)}'

def convert_fhir_to_simplified(fhir_data):
    """Convert FHIR R4 Bundle/Patient to simplified format"""
    
    # Check if it's already in simplified format
    if 'patientId' in fhir_data and 'dateOfBirth' in fhir_data:
        return fhir_data  # Already simplified
    
    # Check if it's a FHIR Bundle
    if fhir_data.get('resourceType') == 'Bundle':
        # Extract Patient resource from Bundle
        patient_resource = None
        conditions = []
        medications = []
        allergies = []
        observations = []
        
        for entry in fhir_data.get('entry', []):
            resource = entry.get('resource', {})
            resource_type = resource.get('resourceType')
            
            if resource_type == 'Patient':
                patient_resource = resource
            elif resource_type == 'Condition':
                conditions.append(resource)
            elif resource_type in ['MedicationRequest', 'MedicationStatement']:
                medications.append(resource)
            elif resource_type == 'AllergyIntolerance':
                allergies.append(resource)
            elif resource_type == 'Observation':
                observations.append(resource)
        
        if not patient_resource:
            raise ValueError("No Patient resource found in Bundle")
        
        # Convert Patient resource
        simplified = convert_fhir_patient(patient_resource)
        
        # Convert related resources
        if conditions:
            simplified['activeConditions'] = [convert_fhir_condition(c) for c in conditions if c.get('clinicalStatus', {}).get('coding', [{}])[0].get('code') == 'active']
            simplified['pastConditions'] = [convert_fhir_condition(c) for c in conditions if c.get('clinicalStatus', {}).get('coding', [{}])[0].get('code') in ['resolved', 'inactive']]
        
        if medications:
            simplified['activeMedications'] = [convert_fhir_medication(m) for m in medications if m.get('status') in ['active', 'intended']]
            simplified['pastMedications'] = [convert_fhir_medication(m) for m in medications if m.get('status') in ['stopped', 'completed']]
        
        if allergies:
            simplified['allergies'] = [convert_fhir_allergy(a) for a in allergies]
        
        return simplified
    
    # Check if it's a single FHIR Patient resource
    elif fhir_data.get('resourceType') == 'Patient':
        return convert_fhir_patient(fhir_data)
    
    else:
        raise ValueError("Unsupported FHIR format. Expected Bundle or Patient resource.")

def convert_fhir_patient(patient):
    """Convert FHIR Patient resource to simplified format"""
    
    # Extract name
    name = "Unknown"
    if patient.get('name') and len(patient['name']) > 0:
        name_obj = patient['name'][0]
        given = ' '.join(name_obj.get('given', []))
        family = name_obj.get('family', '')
        prefix = ' '.join(name_obj.get('prefix', []))
        name = f"{prefix} {given} {family}".strip()
    
    # Extract birthDate
    birth_date = patient.get('birthDate', '')
    
    # Calculate age
    age = 0
    if birth_date:
        from datetime import datetime
        today = datetime.now()
        birth = datetime.fromisoformat(birth_date)
        age = today.year - birth.year
        if today.month < birth.month or (today.month == birth.month and today.day < birth.day):
            age -= 1
    
    # Extract phone
    phone = ""
    if patient.get('telecom'):
        for contact in patient['telecom']:
            if contact.get('system') == 'phone':
                phone = contact.get('value', '')
                break
    
    # Extract email
    email = ""
    if patient.get('telecom'):
        for contact in patient['telecom']:
            if contact.get('system') == 'email':
                email = contact.get('value', '')
                break
    
    # Extract address
    address = ""
    if patient.get('address') and len(patient['address']) > 0:
        addr = patient['address'][0]
        line = ', '.join(addr.get('line', []))
        city = addr.get('city', '')
        state = addr.get('state', '')
        postal = addr.get('postalCode', '')
        address = f"{line}\n{city}, {state} {postal}".strip()
    
    # Generate patient ID from FHIR id
    patient_id = f"PATIENT-{patient.get('id', 'UNKNOWN')[:8]}"
    
    # Build simplified patient
    simplified = {
        'patientId': patient_id,
        'name': name,
        'dateOfBirth': birth_date,
        'age': age,
        'gender': patient.get('gender', 'unknown'),
        'phone': phone,
        'email': email,
        'address': address,
        'emergencyContact': {
            'name': '',
            'relationship': '',
            'phone': ''
        },
        'activeConditions': [],
        'pastConditions': [],
        'activeMedications': [],
        'pastMedications': [],
        'allergies': [],
        'familyHistory': [],
        'socialHistory': {
            'smokingStatus': 'unknown',
            'alcoholUse': '',
            'occupation': '',
            'exerciseHabits': '',
            'diet': ''
        },
        'surgicalHistory': '',
        'immunizations': '',
        'generalNotes': f"Imported from FHIR Bundle. Original FHIR ID: {patient.get('id')}",
        'visits': [],
        'labResults': []
    }
    
    return simplified

def convert_fhir_condition(condition):
    """Convert FHIR Condition to simplified format"""
    
    # Extract condition name
    condition_name = "Unknown Condition"
    if condition.get('code'):
        if condition['code'].get('text'):
            condition_name = condition['code']['text']
        elif condition['code'].get('coding') and len(condition['code']['coding']) > 0:
            condition_name = condition['code']['coding'][0].get('display', 'Unknown')
    
    # Extract dates
    onset_date = ""
    if condition.get('onsetDateTime'):
        onset_date = condition['onsetDateTime'].split('T')[0]
    elif condition.get('recordedDate'):
        onset_date = condition['recordedDate'].split('T')[0]
    
    return {
        'condition': condition_name,
        'diagnosedDate': onset_date,
        'diagnosedBy': '',
        'severity': '',
        'notes': f"FHIR ID: {condition.get('id', 'unknown')}",
        'status': 'active'
    }

def convert_fhir_medication(medication):
    """Convert FHIR Medication to simplified format"""
    
    # Extract medication name
    med_name = "Unknown Medication"
    if medication.get('medicationCodeableConcept'):
        if medication['medicationCodeableConcept'].get('text'):
            med_name = medication['medicationCodeableConcept']['text']
        elif medication['medicationCodeableConcept'].get('coding'):
            med_name = medication['medicationCodeableConcept']['coding'][0].get('display', 'Unknown')
    
    # Extract dosage
    dosage_text = ""
    if medication.get('dosageInstruction') and len(medication['dosageInstruction']) > 0:
        dosage_text = medication['dosageInstruction'][0].get('text', '')
    
    return {
        'medication': med_name,
        'dosage': '',
        'frequency': dosage_text,
        'route': 'oral',
        'startedDate': medication.get('authoredOn', '').split('T')[0] if medication.get('authoredOn') else '',
        'prescribedBy': '',
        'purpose': '',
        'notes': '',
        'status': 'active'
    }

def convert_fhir_allergy(allergy):
    """Convert FHIR AllergyIntolerance to simplified format"""
    
    # Extract allergen
    allergen = "Unknown Allergen"
    if allergy.get('code'):
        if allergy['code'].get('text'):
            allergen = allergy['code']['text']
        elif allergy['code'].get('coding'):
            allergen = allergy['code']['coding'][0].get('display', 'Unknown')
    
    # Extract reaction
    reaction = ""
    severity = "unknown"
    if allergy.get('reaction') and len(allergy['reaction']) > 0:
        reaction_obj = allergy['reaction'][0]
        if reaction_obj.get('manifestation'):
            manifestations = []
            for m in reaction_obj['manifestation']:
                if m.get('text'):
                    manifestations.append(m['text'])
                elif m.get('coding'):
                    manifestations.append(m['coding'][0].get('display', ''))
            reaction = ', '.join(manifestations)
        
        severity = reaction_obj.get('severity', 'unknown')
    
    return {
        'allergen': allergen,
        'reaction': reaction,
        'severity': severity,
        'onsetDate': allergy.get('recordedDate', '').split('T')[0] if allergy.get('recordedDate') else '',
        'notes': f"FHIR ID: {allergy.get('id', 'unknown')}"
    }

# ============================================
# ROUTES
# ============================================

@app.route('/')
def index():
    """Main dashboard"""
    return render_template('index.html')

@app.route('/new-patient')
def new_patient():
    """New patient registration page"""
    return render_template('new-patient.html')

@app.route('/patient-view')
def patient_view():
    """Patient record view page"""
    return render_template('patient-view.html')

@app.route('/new-visit')
def new_visit():
    """New visit form page"""
    return render_template('new-visit.html')

@app.route('/ai-assistant')
def ai_assistant():
    """AI assistant page"""
    return render_template('ai-assistant.html')

@app.route('/settings')
def settings():
    """Settings page for LLM configuration"""
    return render_template('settings.html')

# ============================================
# API ENDPOINTS
# ============================================

@app.route('/api/generate-patient-id', methods=['GET'])
def api_generate_patient_id():
    """Generate a new patient ID"""
    return jsonify({'patientId': generate_patient_id()})

@app.route('/api/patients', methods=['GET'])
def api_get_patients():
    """Get all patients"""
    patients = load_patients()
    return jsonify(patients)

@app.route('/api/patients/<patient_id>', methods=['GET'])
def api_get_patient(patient_id):
    """Get a specific patient by ID"""
    patients = load_patients()
    patient = next((p for p in patients if p['patientId'] == patient_id), None)
    
    if patient:
        return jsonify(patient)
    else:
        return jsonify({'error': 'Patient not found'}), 404

@app.route('/api/patients/search', methods=['POST'])
def api_search_patients():
    """Search patients by ID or name"""
    data = request.json
    search_term = data.get('searchTerm', '').lower()
    
    patients = load_patients()
    results = [
        p for p in patients 
        if search_term in p['patientId'].lower() or 
           search_term in p['name'].lower()
    ]
    
    return jsonify(results)

@app.route('/api/patients', methods=['POST'])
def api_create_patient():
    """Create a new patient"""
    patient_data = request.json
    
    # Validate required fields
    required_fields = ['patientId', 'name', 'dateOfBirth', 'gender']
    for field in required_fields:
        if field not in patient_data:
            return jsonify({'error': f'Missing required field: {field}'}), 400
    
    # Load existing patients
    patients = load_patients()
    
    # Check if patient ID already exists
    if any(p['patientId'] == patient_data['patientId'] for p in patients):
        return jsonify({'error': 'Patient ID already exists'}), 400
    
    # Add patient
    patients.append(patient_data)
    save_patients(patients)
    
    return jsonify({'success': True, 'patientId': patient_data['patientId']})

@app.route('/api/patients/<patient_id>', methods=['PUT'])
def api_update_patient(patient_id):
    """Update an existing patient"""
    patient_data = request.json
    patients = load_patients()
    
    # Find patient index
    index = next((i for i, p in enumerate(patients) if p['patientId'] == patient_id), None)
    
    if index is None:
        return jsonify({'error': 'Patient not found'}), 404
    
    # Update patient
    patients[index] = patient_data
    save_patients(patients)
    
    return jsonify({'success': True})

@app.route('/api/patients/<patient_id>/visits', methods=['POST'])
def api_add_visit(patient_id):
    """Add a visit to a patient"""
    visit_data = request.json
    patients = load_patients()
    
    # Find patient
    patient = next((p for p in patients if p['patientId'] == patient_id), None)
    
    if not patient:
        return jsonify({'error': 'Patient not found'}), 404
    
    # Ensure visits array exists
    if 'visits' not in patient:
        patient['visits'] = []
    
    # Generate visit ID
    visit_data['visitId'] = generate_visit_id(patient_id)
    
    # Add visit
    patient['visits'].append(visit_data)
    
    # Save patients
    save_patients(patients)
    
    return jsonify({'success': True, 'visitId': visit_data['visitId']})

@app.route('/api/patients/<patient_id>/visits/<visit_id>/report', methods=['GET'])
def api_generate_visit_report(patient_id, visit_id):
    """Generate and download Word report for a specific visit"""
    patients = load_patients()
    patient = next((p for p in patients if p['patientId'] == patient_id), None)
    
    if not patient:
        return jsonify({'error': 'Patient not found'}), 404
    
    # Find the specific visit
    visit = None
    if patient.get('visits'):
        visit = next((v for v in patient['visits'] if v.get('visitId') == visit_id), None)
    
    if not visit:
        return jsonify({'error': 'Visit not found'}), 404
    
    try:
        # Generate Word document
        doc = generate_visit_report_docx(patient, visit)
        
        # Save to BytesIO object
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Create filename
        filename = f"Visit_Report_{patient_id}_{visit_id}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': f'Error generating report: {str(e)}'}), 500

@app.route('/api/patients/<patient_id>/full-report', methods=['GET'])
def api_generate_patient_report(patient_id):
    """Generate and download complete patient medical record as Word document"""
    patients = load_patients()
    patient = next((p for p in patients if p['patientId'] == patient_id), None)
    
    if not patient:
        return jsonify({'error': 'Patient not found'}), 404
    
    try:
        # Generate Word document
        doc = generate_patient_report_docx(patient)
        
        # Save to BytesIO object
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Create filename
        patient_name = patient.get('name', 'Unknown').replace(' ', '_')
        filename = f"Complete_Medical_Record_{patient_id}_{patient_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': f'Error generating report: {str(e)}'}), 500

@app.route('/api/patients/<patient_id>', methods=['DELETE'])
def api_delete_patient(patient_id):
    """Delete a patient"""
    patients = load_patients()
    
    # Find and remove patient
    patients = [p for p in patients if p['patientId'] != patient_id]
    save_patients(patients)
    
    return jsonify({'success': True})

@app.route('/api/patients/<patient_id>/export', methods=['GET'])
def api_export_patient(patient_id):
    """Export patient data as JSON"""
    patients = load_patients()
    patient = next((p for p in patients if p['patientId'] == patient_id), None)
    
    if not patient:
        return jsonify({'error': 'Patient not found'}), 404
    
    return jsonify(patient)

@app.route('/api/patients/import', methods=['POST'])
def api_import_patient():
    """Import patient data from JSON (supports FHIR R4 or simplified format)"""
    patient_data = request.json
    
    try:
        # Convert FHIR R4 to simplified format if needed
        patient_data = convert_fhir_to_simplified(patient_data)
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    
    # Validate required fields (after conversion)
    required_fields = ['patientId', 'name', 'dateOfBirth', 'gender']
    for field in required_fields:
        if field not in patient_data:
            return jsonify({'error': f'Missing required field: {field}'}), 400
    
    patients = load_patients()
    
    # Check if patient ID already exists
    existing_patient = next((p for p in patients if p['patientId'] == patient_data['patientId']), None)
    
    if existing_patient:
        # Ask if want to replace
        return jsonify({'error': 'Patient ID already exists', 'exists': True, 'patientId': patient_data['patientId']}), 409
    
    # Add patient
    patients.append(patient_data)
    save_patients(patients)
    
    return jsonify({'success': True, 'patientId': patient_data['patientId']})

@app.route('/api/patients/import/replace', methods=['POST'])
def api_import_patient_replace():
    """Import patient data and replace if exists (supports FHIR R4 or simplified format)"""
    patient_data = request.json
    
    try:
        # Convert FHIR R4 to simplified format if needed
        patient_data = convert_fhir_to_simplified(patient_data)
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    
    patients = load_patients()
    
    # Find and replace or add
    index = next((i for i, p in enumerate(patients) if p['patientId'] == patient_data['patientId']), None)
    
    if index is not None:
        patients[index] = patient_data
    else:
        patients.append(patient_data)
    
    save_patients(patients)
    return jsonify({'success': True, 'patientId': patient_data['patientId']})

@app.route('/api/settings', methods=['GET'])
def api_get_settings():
    """Get current LLM settings"""
    settings = load_settings()
    # Don't send full model path for security, just filename
    if settings.get('model_path'):
        settings['model_filename'] = os.path.basename(settings['model_path'])
    return jsonify(settings)

@app.route('/api/settings', methods=['POST'])
def api_save_settings():
    """Save Ollama settings"""
    try:
        new_settings = request.json
        save_settings(new_settings)
        
        return jsonify({'success': True})
    except Exception as e:
        print(f"Error saving settings: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/settings/test-model', methods=['POST'])
def api_test_model():
    """Test Ollama model"""
    try:
        settings = request.json
        
        print("Testing Ollama...")
        try:
            import ollama
            
            model_name = settings.get('ollama_model', 'qwen2.5:4b')
            print(f"Testing Ollama model: {model_name}")
            
            # Simple test query
            response = ollama.chat(
                model=model_name,
                messages=[
                    {'role': 'user', 'content': 'Say hello in one sentence.'}
                ],
                options={
                    'num_predict': 50
                }
            )
            
            return jsonify({
                'message': f'✓ Ollama Test Successful!\n\nModel: {model_name}\nStatus: Running\nResponse: {response["message"]["content"][:100]}'
            })
            
        except ImportError:
            return jsonify({'error': 'Ollama not installed.\n\nInstall: pip install ollama\nDownload: https://ollama.com/download'}), 500
        except Exception as e:
            error_msg = str(e)
            if 'connect' in error_msg.lower():
                return jsonify({'error': 'Cannot connect to Ollama.\n\nMake sure Ollama is running (check system tray or run: ollama serve)'}), 500
            elif 'not found' in error_msg.lower():
                return jsonify({'error': f'Model not found: {model_name}\n\nDownload it with: ollama pull {model_name}'}), 500
            else:
                return jsonify({'error': f'Ollama error: {error_msg}'}), 500
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Model test error: {error_details}")
        return jsonify({'error': f'Test failed: {str(e)}'}), 500

@app.route('/api/ollama/unload', methods=['POST'])
def api_ollama_unload():
    """Unload Ollama model from memory"""
    try:
        import ollama
        import subprocess
        
        # Run ollama ps to see loaded models
        result = subprocess.run(['ollama', 'ps'], capture_output=True, text=True)
        
        if result.returncode != 0:
            return jsonify({'error': 'Could not check Ollama status'}), 500
        
        # Check if any models are loaded
        output = result.stdout
        lines = output.strip().split('\n')
        
        if len(lines) <= 1:  # Only header, no models loaded
            return jsonify({'message': '✓ No models currently loaded in memory'}), 200
        
        # Models are loaded - unload them
        # We'll use a shell command to stop ollama (which unloads models)
        # Then restart it
        try:
            # On Windows, kill the ollama process to free memory
            # On Linux/Mac, use ollama stop
            import platform
            
            if platform.system() == 'Windows':
                # Use PowerShell to restart Ollama service
                subprocess.run(['powershell', '-Command', 'Stop-Process', '-Name', 'ollama', '-Force'], 
                             capture_output=True)
                # Wait a moment
                import time
                time.sleep(1)
                # Ollama will auto-restart as a service on Windows
                
                return jsonify({
                    'message': '✓ Ollama restarted - all models unloaded from memory!\n\nMemory has been freed. Next query will reload the model.'
                }), 200
            else:
                # On Linux/Mac, we can use ollama's built-in command
                subprocess.run(['pkill', 'ollama'], capture_output=True)
                import time
                time.sleep(1)
                
                return jsonify({
                    'message': '✓ Ollama stopped - all models unloaded from memory!\n\nMemory has been freed. Ollama will restart on next query.'
                }), 200
                
        except Exception as e:
            return jsonify({
                'error': f'Could not unload model: {str(e)}\n\nYou can manually stop Ollama from system tray or task manager.'
            }), 500
            
    except ImportError:
        return jsonify({'error': 'Ollama not installed'}), 500
    except Exception as e:
        return jsonify({'error': f'Error: {str(e)}'}), 500

@app.route('/api/ai-query', methods=['POST'])
def api_ai_query():
    """Query Ollama AI with patient context"""
    data = request.json
    patient_id = data.get('patientId')
    question = data.get('question')
    
    if not patient_id or not question:
        return jsonify({'error': 'Missing patientId or question'}), 400
    
    # Get patient data
    patients = load_patients()
    patient = next((p for p in patients if p['patientId'] == patient_id), None)
    
    if not patient:
        return jsonify({'error': 'Patient not found'}), 404
    
    # Build context
    context = build_patient_context(patient)
    
    # Get Ollama settings
    settings = load_settings()
    
    try:
        import ollama
        
        system_prompt = settings.get('system_prompt', 'You are a clinical AI assistant.')
        full_prompt = f"{context}\n\nDoctor's Question: {question}"
        
        response = ollama.chat(
            model=settings.get('ollama_model', 'qwen2.5:4b'),
            messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': full_prompt}
            ],
            options={
                'temperature': settings.get('temperature', 0.7),
                'num_predict': settings.get('max_tokens', 2048)
            }
        )
        
        return jsonify({
            'response': response['message']['content'],
            'source': 'ollama'
        })
        
    except ImportError:
        return jsonify({'error': 'Ollama not installed.\n\nInstall: pip install ollama\nDownload: https://ollama.com/download'}), 500
    except Exception as e:
        error_msg = str(e)
        if 'connect' in error_msg.lower():
            return jsonify({'error': 'Cannot connect to Ollama.\n\nMake sure Ollama is running.'}), 500
        elif 'not found' in error_msg.lower():
            model = settings.get('ollama_model', 'qwen2.5:4b')
            return jsonify({'error': f'Model not found: {model}\n\nDownload: ollama pull {model}'}), 500
        else:
            return jsonify({'error': f'Ollama error: {error_msg}'}), 500

def build_patient_context(patient):
    """Build formatted patient context for AI"""
    context = f"PATIENT RECORD\n"
    context += f"=" * 50 + "\n\n"
    
    # Demographics
    context += f"Patient ID: {patient.get('patientId', 'N/A')}\n"
    context += f"Name: {patient.get('name', 'N/A')}\n"
    context += f"Age: {patient.get('age', 'N/A')} years old\n"
    context += f"Gender: {patient.get('gender', 'N/A')}\n"
    context += f"Date of Birth: {patient.get('dateOfBirth', 'N/A')}\n"
    context += f"\n"
    
    # Active Conditions
    if patient.get('activeConditions'):
        context += "ACTIVE CONDITIONS:\n"
        for cond in patient['activeConditions']:
            context += f"  - {cond.get('condition', 'N/A')} (diagnosed {cond.get('diagnosedDate', 'N/A')})\n"
            if cond.get('notes'):
                context += f"    Notes: {cond['notes']}\n"
        context += "\n"
    
    # Past Conditions
    if patient.get('pastConditions'):
        context += "PAST CONDITIONS (RESOLVED):\n"
        for cond in patient['pastConditions']:
            context += f"  - {cond.get('condition', 'N/A')} ({cond.get('diagnosedDate', 'N/A')} to {cond.get('resolvedDate', 'N/A')})\n"
        context += "\n"
    
    # Current Medications
    if patient.get('activeMedications'):
        context += "CURRENT MEDICATIONS:\n"
        for med in patient['activeMedications']:
            context += f"  - {med.get('medication', 'N/A')} {med.get('dosage', '')} {med.get('frequency', '')}\n"
            if med.get('purpose'):
                context += f"    Purpose: {med['purpose']}\n"
        context += "\n"
    
    # Past Medications
    if patient.get('pastMedications'):
        context += "DISCONTINUED MEDICATIONS:\n"
        for med in patient['pastMedications']:
            context += f"  - {med.get('medication', 'N/A')} (stopped {med.get('stoppedDate', 'N/A')})\n"
            if med.get('reasonStopped'):
                context += f"    Reason: {med['reasonStopped']}\n"
        context += "\n"
    
    # Allergies
    if patient.get('allergies'):
        context += "ALLERGIES:\n"
        for allergy in patient['allergies']:
            context += f"  - {allergy.get('allergen', 'N/A')} - {allergy.get('reaction', 'N/A')} ({allergy.get('severity', 'N/A')} severity)\n"
        context += "\n"
    
    # Family History
    if patient.get('familyHistory'):
        context += "FAMILY HISTORY:\n"
        for family in patient['familyHistory']:
            context += f"  {family.get('relative', 'N/A')} ({family.get('age', 'N/A')}, {family.get('livingStatus', 'N/A')}):\n"
            if family.get('conditions'):
                for cond in family['conditions']:
                    # Handle both string and dict formats
                    if isinstance(cond, str):
                        context += f"    - {cond}\n"
                    else:
                        context += f"    - {cond.get('condition', 'N/A')} (onset age {cond.get('ageOfOnset', 'N/A')})\n"
        context += "\n"
    
    # Social History
    if patient.get('socialHistory'):
        sh = patient['socialHistory']
        context += "SOCIAL HISTORY:\n"
        context += f"  Smoking: {sh.get('smokingStatus', 'N/A')}\n"
        context += f"  Alcohol: {sh.get('alcoholUse', 'N/A')}\n"
        context += f"  Occupation: {sh.get('occupation', 'N/A')}\n"
        context += f"  Exercise: {sh.get('exerciseHabits', 'N/A')}\n"
        context += "\n"
    
    # Recent Visits
    if patient.get('visits'):
        context += "VISIT HISTORY:\n"
        recent_visits = patient['visits'][-5:]  # Last 5 visits
        for visit in recent_visits:
            context += f"\n  Visit Date: {visit.get('date', 'N/A')}\n"
            context += f"  Type: {visit.get('type', 'N/A')}\n"
            context += f"  Chief Complaint: {visit.get('chiefComplaint', 'N/A')}\n"
            
            if visit.get('vitals'):
                vitals = visit['vitals']
                context += f"  Vitals: BP {vitals.get('bloodPressure', 'N/A')}, HR {vitals.get('heartRate', 'N/A')}, Temp {vitals.get('temperature', 'N/A')}\n"
            
            context += f"  Symptoms: {visit.get('symptoms', 'N/A')}\n"
            context += f"  Assessment: {visit.get('assessment', 'N/A')}\n"
            context += f"  Plan: {visit.get('plan', 'N/A')}\n"
        context += "\n"
    
    # Lab Results
    if patient.get('labResults'):
        context += "RECENT LAB RESULTS:\n"
        for lab in patient['labResults'][-5:]:  # Last 5 labs
            context += f"  {lab.get('testName', 'N/A')} ({lab.get('date', 'N/A')}): {lab.get('result', 'N/A')}\n"
        context += "\n"
    
    return context

def generate_visit_report_docx(patient, visit):
    """Generate Word document for a single visit"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add clinic header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(CLINIC_NAME)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(102, 126, 234)
    
    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header2.add_run(f'{CLINIC_ADDRESS}\n{CLINIC_PHONE}')
    run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('VISIT REPORT')
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph('_' * 80)
    
    # Patient Information
    doc.add_heading('Patient Information', level=2)
    patient_info = doc.add_paragraph()
    patient_info.add_run(f"Name: ").bold = True
    patient_info.add_run(f"{patient.get('name', 'N/A')}\n")
    patient_info.add_run(f"Patient ID: ").bold = True
    patient_info.add_run(f"{patient.get('patientId', 'N/A')}\n")
    patient_info.add_run(f"Date of Birth: ").bold = True
    patient_info.add_run(f"{patient.get('dateOfBirth', 'N/A')} (Age: {patient.get('age', 'N/A')})\n")
    patient_info.add_run(f"Gender: ").bold = True
    patient_info.add_run(f"{patient.get('gender', 'N/A').title()}\n")
    
    # Visit Information
    doc.add_heading('Visit Information', level=2)
    visit_info = doc.add_paragraph()
    visit_info.add_run(f"Visit ID: ").bold = True
    visit_info.add_run(f"{visit.get('visitId', 'N/A')}\n")
    visit_info.add_run(f"Date: ").bold = True
    visit_info.add_run(f"{visit.get('date', 'N/A')}\n")
    visit_info.add_run(f"Type: ").bold = True
    visit_info.add_run(f"{visit.get('type', 'N/A')}\n")
    visit_info.add_run(f"Provider: ").bold = True
    visit_info.add_run(f"{visit.get('seenBy', 'N/A')}\n")
    
    # Chief Complaint
    doc.add_heading('Chief Complaint', level=2)
    doc.add_paragraph(visit.get('chiefComplaint', 'N/A'))
    
    # Vital Signs
    if visit.get('vitals'):
        doc.add_heading('Vital Signs', level=2)
        vitals = visit['vitals']
        
        # Create table for vitals
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        vitals_data = [
            ('Blood Pressure', vitals.get('bloodPressure', 'N/A')),
            ('Heart Rate', f"{vitals.get('heartRate', 'N/A')} bpm"),
            ('Temperature', f"{vitals.get('temperature', 'N/A')} °F"),
            ('Weight', f"{vitals.get('weight', 'N/A')} lbs (BMI: {vitals.get('bmi', 'N/A')})"),
            ('Oxygen Saturation', f"{vitals.get('oxygenSat', 'N/A')}%")
        ]
        
        for i, (label, value) in enumerate(vitals_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            table.rows[i].cells[1].text = value
    
    # Symptoms
    doc.add_heading('Symptoms', level=2)
    doc.add_paragraph(visit.get('symptoms', 'N/A'))
    
    # Physical Examination
    doc.add_heading('Physical Examination', level=2)
    doc.add_paragraph(visit.get('examination', 'N/A'))
    
    # Assessment/Diagnosis
    doc.add_heading('Assessment & Diagnosis', level=2)
    doc.add_paragraph(visit.get('assessment', 'N/A'))
    
    # Treatment Plan
    doc.add_heading('Treatment Plan', level=2)
    doc.add_paragraph(visit.get('plan', 'N/A'))
    
    # Prescriptions
    if visit.get('prescriptions') and len(visit['prescriptions']) > 0:
        doc.add_heading('Prescriptions', level=2)
        for rx in visit['prescriptions']:
            if rx.get('medication'):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{rx.get('medication')} ").bold = True
                p.add_run(f"{rx.get('dosage', '')} - {rx.get('frequency', '')}")
                if rx.get('quantity'):
                    p.add_run(f"\n   Quantity: {rx['quantity']}, Refills: {rx.get('refills', '0')}")
    
    # Labs Ordered
    if visit.get('labsOrdered'):
        doc.add_heading('Laboratory Tests Ordered', level=2)
        doc.add_paragraph(visit.get('labsOrdered'))
    
    # Follow-up
    if visit.get('followUp'):
        doc.add_heading('Follow-up', level=2)
        doc.add_paragraph(f"Scheduled for: {visit.get('followUp')}")
    
    # Notes
    if visit.get('notes'):
        doc.add_heading('Additional Notes', level=2)
        doc.add_paragraph(visit.get('notes'))
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'\nReport Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n')
    run.font.size = Pt(9)
    run.font.italic = True
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run('CONFIDENTIAL MEDICAL RECORD\nFor authorized use only')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    return doc

def generate_patient_report_docx(patient):
    """Generate comprehensive Word document for complete patient record"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add clinic header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(CLINIC_NAME)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(102, 126, 234)
    
    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header2.add_run(f'{CLINIC_ADDRESS}\n{CLINIC_PHONE}')
    run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('COMPREHENSIVE MEDICAL RECORD')
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph('_' * 80)
    
    # Patient Demographics
    doc.add_heading('Patient Demographics', level=1)
    demo_table = doc.add_table(rows=6, cols=2)
    demo_table.style = 'Light Grid Accent 1'
    
    demo_data = [
        ('Full Name', patient.get('name', 'N/A')),
        ('Patient ID', patient.get('patientId', 'N/A')),
        ('Date of Birth', f"{patient.get('dateOfBirth', 'N/A')} (Age: {patient.get('age', 'N/A')})"),
        ('Gender', patient.get('gender', 'N/A').title()),
        ('Phone', patient.get('phone', 'N/A')),
        ('Email', patient.get('email', 'N/A'))
    ]
    
    for i, (label, value) in enumerate(demo_data):
        demo_table.rows[i].cells[0].text = label
        demo_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        demo_table.rows[i].cells[1].text = value
    
    # Address
    if patient.get('address'):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Address: ').bold = True
        p.add_run(patient.get('address'))
    
    # Emergency Contact
    if patient.get('emergencyContact'):
        ec = patient['emergencyContact']
        if ec.get('name'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Emergency Contact: ').bold = True
            p.add_run(f"{ec.get('name')} ({ec.get('relationship', 'N/A')}) - {ec.get('phone', 'N/A')}")
    
    # Active Conditions
    doc.add_heading('Active Medical Conditions', level=1)
    if patient.get('activeConditions') and len(patient['activeConditions']) > 0:
        for cond in patient['activeConditions']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{cond.get('condition', 'N/A')}").bold = True
            details = f"\n   Diagnosed: {cond.get('diagnosedDate', 'N/A')}"
            if cond.get('diagnosedBy'):
                details += f" by {cond['diagnosedBy']}"
            if cond.get('severity'):
                details += f" | Severity: {cond['severity']}"
            if cond.get('notes'):
                details += f"\n   Notes: {cond['notes']}"
            p.add_run(details)
    else:
        doc.add_paragraph('No active conditions recorded.')
    
    # Past Conditions
    if patient.get('pastConditions') and len(patient['pastConditions']) > 0:
        doc.add_heading('Past Medical Conditions (Resolved)', level=2)
        for cond in patient['pastConditions']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{cond.get('condition', 'N/A')}")
            p.add_run(f" ({cond.get('diagnosedDate', 'N/A')} to {cond.get('resolvedDate', 'N/A')})")
    
    # Current Medications
    doc.add_heading('Current Medications', level=1)
    if patient.get('activeMedications') and len(patient['activeMedications']) > 0:
        for med in patient['activeMedications']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{med.get('medication', 'N/A')} ").bold = True
            p.add_run(f"{med.get('dosage', '')} - {med.get('frequency', '')} ({med.get('route', 'oral')})")
            details = f"\n   Started: {med.get('startedDate', 'N/A')}"
            if med.get('purpose'):
                details += f" | Purpose: {med['purpose']}"
            if med.get('prescribedBy'):
                details += f"\n   Prescribed by: {med['prescribedBy']}"
            p.add_run(details)
    else:
        doc.add_paragraph('No current medications.')
    
    # Allergies (HIGHLIGHTED)
    doc.add_heading('⚠️ ALLERGIES & ADVERSE REACTIONS', level=1)
    if patient.get('allergies') and len(patient['allergies']) > 0:
        allergy_table = doc.add_table(rows=len(patient['allergies']) + 1, cols=3)
        allergy_table.style = 'Medium Shading 1 Accent 2'
        
        # Header row
        header_cells = allergy_table.rows[0].cells
        header_cells[0].text = 'Allergen'
        header_cells[1].text = 'Reaction'
        header_cells[2].text = 'Severity'
        
        for i, cell in enumerate(header_cells):
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Data rows
        for i, allergy in enumerate(patient['allergies']):
            row = allergy_table.rows[i + 1]
            row.cells[0].text = allergy.get('allergen', 'N/A')
            row.cells[1].text = allergy.get('reaction', 'N/A')
            row.cells[2].text = allergy.get('severity', 'N/A').upper()
    else:
        doc.add_paragraph('No known allergies.')
    
    # Family History
    if patient.get('familyHistory') and len(patient['familyHistory']) > 0:
        doc.add_heading('Family Medical History', level=1)
        for family in patient['familyHistory']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{family.get('relative', 'N/A')} ").bold = True
            p.add_run(f"({family.get('age', 'N/A')} years, {family.get('livingStatus', 'N/A')})")
            if family.get('conditions'):
                p.add_run('\n   Conditions: ' + family.get('conditions', ''))
            if family.get('notes'):
                p.add_run(f"\n   Notes: {family['notes']}")
    
    # Social History
    if patient.get('socialHistory'):
        doc.add_heading('Social History', level=1)
        sh = patient['socialHistory']
        social_table = doc.add_table(rows=5, cols=2)
        social_table.style = 'Light Grid Accent 1'
        
        social_data = [
            ('Smoking Status', sh.get('smokingStatus', 'N/A')),
            ('Alcohol Use', sh.get('alcoholUse', 'N/A')),
            ('Occupation', sh.get('occupation', 'N/A')),
            ('Exercise Habits', sh.get('exerciseHabits', 'N/A')),
            ('Diet', sh.get('diet', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(social_data):
            social_table.rows[i].cells[0].text = label
            social_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            social_table.rows[i].cells[1].text = value
    
    # Visit History
    if patient.get('visits') and len(patient['visits']) > 0:
        doc.add_heading('Visit History', level=1)
        doc.add_paragraph(f'Total Visits: {len(patient["visits"])}')
        
        # Show last 5 visits in detail
        recent_visits = patient['visits'][-5:]
        recent_visits.reverse()  # Most recent first
        
        for visit in recent_visits:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run(f"Visit Date: {visit.get('date', 'N/A')}").bold = True
            p.add_run(f" | Type: {visit.get('type', 'N/A')}")
            
            p2 = doc.add_paragraph(style='List Bullet 2')
            p2.add_run('Chief Complaint: ').bold = True
            p2.add_run(visit.get('chiefComplaint', 'N/A'))
            
            if visit.get('vitals'):
                v = visit['vitals']
                p3 = doc.add_paragraph(style='List Bullet 2')
                p3.add_run('Vitals: ')
                p3.add_run(f"BP {v.get('bloodPressure', 'N/A')}, HR {v.get('heartRate', 'N/A')}, Temp {v.get('temperature', 'N/A')}")
            
            p4 = doc.add_paragraph(style='List Bullet 2')
            p4.add_run('Assessment: ').bold = True
            p4.add_run(visit.get('assessment', 'N/A'))
            
            p5 = doc.add_paragraph(style='List Bullet 2')
            p5.add_run('Plan: ').bold = True
            p5.add_run(visit.get('plan', 'N/A'))
    
    # Additional Information
    if patient.get('surgicalHistory'):
        doc.add_heading('Surgical History', level=1)
        doc.add_paragraph(patient.get('surgicalHistory'))
    
    if patient.get('immunizations'):
        doc.add_heading('Immunizations', level=1)
        doc.add_paragraph(patient.get('immunizations'))
    
    if patient.get('generalNotes'):
        doc.add_heading('General Notes', level=1)
        doc.add_paragraph(patient.get('generalNotes'))
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'\nComplete Medical Record Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n')
    run.font.size = Pt(9)
    run.font.italic = True
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run('CONFIDENTIAL MEDICAL RECORD\nFor authorized use only\nThis document contains protected health information')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    return doc

if __name__ == '__main__':
    print("\n" + "="*50)
    print("Healthcare Management System")
    print("="*50)
    print("\nStarting server...")
    print("Access the application at: http://localhost:5000")
    print("\nPress Ctrl+C to stop the server\n")
    app.run(debug=True, host='0.0.0.0', port=5000)
